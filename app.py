import re
import io
import asyncio
from typing import List, Dict, Optional, Tuple
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
import httpx
from rapidfuzz import fuzz

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Pt

# -----------------------------------------------------------------------------
# APP INIT + CORS
# -----------------------------------------------------------------------------
app = FastAPI(
    @app.get("/")
def root():
    return {"ok": True, "service": "AI Detector PRO"}

@app.get("/health")
def health():
    return {"ok": True}

    title="AI Detector PRO",
    version="3.2.0",
    openapi_url="/openapi.json",
    docs_url="/docs",
    redoc_url="/redoc",
)

@app.get("/health")
def health():
    return {"ok": True}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_headers=["*"],
    allow_methods=["*"],
    expose_headers=["*"],
)

# -----------------------------------------------------------------------------
# UTILS
# -----------------------------------------------------------------------------
_SENTENCE_RE = re.compile(r'(?us)([^.!?]+[.!?])')
_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"\s+")

HTTP_LIMITS = httpx.Limits(max_connections=5, max_keepalive_connections=5)
HTTP_TIMEOUT = httpx.Timeout(connect=5, read=8, write=8, pool=5)

def clean_text(text: str) -> str:
    text = _TAG_RE.sub(" ", text)
    return _WS_RE.sub(" ", text).strip()

def split_sentences(text: str) -> List[str]:
    parts = _SENTENCE_RE.findall(text) or [text]
    return [p.strip() for p in parts if p.strip()]

def stitch(xs: List[str]):
    return " ".join(x.strip() for x in xs if x.strip())

def ai_score(text: str) -> int:
    sents = split_sentences(text)
    long_count = sum(len(s.split()) >= 18 for s in sents)
    return min(100, 15 + long_count * 12)

async def fetch(url: str, client: httpx.AsyncClient) -> str:
    try:
        r = await client.get(url)
        if r.status_code == 200:
            return clean_text(r.text)
    except:
        pass
    return ""

async def build_corpus(texts: List[str], urls: List[str]):
    corpus = [(f"text:{i}", clean_text(t)) for i, t in enumerate(texts or [])]
    if urls:
        urls = urls[:5]
        async with httpx.AsyncClient(
            follow_redirects=True,
            limits=HTTP_LIMITS,
            timeout=HTTP_TIMEOUT,
        ) as client:
            results = await asyncio.gather(*[fetch(u, client) for u in urls])
        for u, body in zip(urls, results):
            if body:
                corpus.append((u, body))
    return corpus

def similarity(s: str, body: str):
    best, snip = 0, ""
    for cs in split_sentences(body):
        sc = max(fuzz.token_set_ratio(s, cs), fuzz.partial_ratio(s, cs))
        if sc > best:
            best, snip = sc, cs
    return best, snip

def reason_ai(s: str):
    long = len(s.split()) >= 22
    passive = any(x in s.lower() for x in ["is being", "was", "were", "been"])
    lowvar = len(set(s.lower().split())) < len(s.split()) * 0.65
    r = []
    if long: r.append("very long")
    if passive: r.append("passive voice")
    if lowvar: r.append("low variety")
    return ", ".join(r) or "generic style"

# -----------------------------------------------------------------------------
# MODELS
# -----------------------------------------------------------------------------
class AnalyzeIn(BaseModel):
    text: str
    source_texts: Optional[List[str]] = None
    source_urls: Optional[List[str]] = None
    plag_threshold: int = 85

class SentOut(BaseModel):
    index: int
    text: str
    label: str
    confidence: str
    plag_score: Optional[int] = None
    source: Optional[str] = None
    snippet: Optional[str] = None
    citation_index: Optional[int] = None
    reason: Optional[str] = None
    fix: Optional[str] = None

class AnalyzeOut(BaseModel):
    summary: dict
    counts: dict
    sentences: List[SentOut]
    citations: List[str]
    markdown_report: str

# -----------------------------------------------------------------------------
# CORE ANALYSIS
# -----------------------------------------------------------------------------
async def core(req: AnalyzeIn):
    text = req.text.strip()
    sents = split_sentences(text)
    ai_over = ai_score(text)
    corpus = await build_corpus(req.source_texts or [], req.source_urls or [])

    cite_map = {}
    citations = []
    out = []
    ai_n = pl_n = 0

    for i, s in enumerate(sents):
        best, src, snip = 0, None, None
        for u, body in corpus:
            sc, mp = similarity(s, body)
            if sc > best:
                best, src, snip = sc, u, mp

        label = "human"
        if best >= req.plag_threshold:
            label = "plag"
        elif len(s.split()) >= 18 or ai_over >= 70:
            label = "ai"

        if label == "ai": ai_n += 1
        if label == "plag": pl_n += 1

        conf = "High" if (label == "plag" and best > 92) or (label == "ai" and ai_over > 80) else ("Medium" if label != "human" else "Low")

        cidx = None
        if label == "plag" and src:
            if src not in cite_map:
                cite_map[src] = len(citations) + 1
                citations.append(src)
            cidx = cite_map[src]

        out.append(SentOut(
            index=i,
            text=s,
            label=label,
            confidence=conf,
            plag_score=best if best else None,
            source=src,
            snippet=snip,
            citation_index=cidx,
            reason=reason_ai(s) if label == "ai" else ("high similarity" if label == "plag" else None),
            fix="Rewrite with personal detail" if label=="ai" else ("Rephrase & cite source" if label=="plag" else None),
        ))

    total = len(out)
    plag_pct = min(100, int(100 * pl_n / max(1, total)))
    summary = {
        "ai_score": ai_over,
        "plagiarism_score": plag_pct,
        "human_score": max(0, 100 - max(ai_over, plag_pct)),
    }
    counts = {"total": total, "ai": ai_n, "plag": pl_n, "human": total-ai_n-pl_n}

    md = [
        "## AI Text Detection Report",
        f"- **AI-likeness:** {summary['ai_score']}%",
        f"- **Plagiarism risk:** {summary['plagiarism_score']}%",
        f"- **Human-likeness:** {summary['human_score']}%",
        "⚠️ Results are indicators. Verify with context."
    ]

    if ai_n:
        md.append("\n### AI-like sentences")
        for s in out:
            if s.label == "ai":
                md.append(f"- {s.text} (Reason: {s.reason})")
    if pl_n:
        md.append("\n### Plagiarism-risk sentences")
        for s in out:
            if s.label == "plag":
                md.append(f"- {s.text} (Similarity: {s.plag_score}%)")

    return AnalyzeOut(
        summary=summary,
        counts=counts,
        sentences=out,
        citations=citations,
        markdown_report="\n".join(md)
    )

# -----------------------------------------------------------------------------
# ENDPOINTS
# -----------------------------------------------------------------------------
@app.post("/analyze", response_model=AnalyzeOut)
async def analyze(req: AnalyzeIn):
    return JSONResponse((await core(req)).model_dump())

class RewriteIn(BaseModel):
    sentence: str
    style: str = "humanize"

@app.post("/rewrite")
async def rewrite(req: RewriteIn):
    s = req.sentence.strip()
    return {
        "original": req.sentence,
        "rewritten": s + " (more human detail)",
        "style": req.style,
    }

class BatchIn(BaseModel):
    text: str

@app.post("/rewrite-batch")
async def rewrite_batch(req: BatchIn):
    sents = split_sentences(req.text)
    return {"rewritten": stitch([s + " (humanized)" for s in sents])}

class PdfIn(BaseModel):
    text: str
    filename: str = "detector_report.pdf"

@app.post("/report.pdf")
async def report_pdf(req: PdfIn):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(40,760,"AI Text Detection Report")
    c.drawString(40,740,"Download complete ✅")
    c.save()
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition":f'attachment; filename="{req.filename}"'}
    )

class DocxIn(BaseModel):
    text: str
    filename: str = "detector_report.docx"

@app.post("/report.docx")
async def report_docx(req: DocxIn):
    doc = Document()
    doc.add_heading('AI Report', 0)
    doc.add_paragraph("The report downloaded successfully ✅")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":f'attachment; filename="{req.filename}"'},
    )
